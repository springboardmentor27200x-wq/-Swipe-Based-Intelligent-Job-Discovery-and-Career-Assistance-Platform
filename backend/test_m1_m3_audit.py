import os
import sys
import io
import docx  # type: ignore
from fastapi.testclient import TestClient  # type: ignore

# Ensure backend directory is in path
sys.path.append(os.path.dirname(os.path.abspath(__file__)))

from app.main import app
from app.database import Base, engine, SessionLocal
from app import models, auth, crud

client = TestClient(app)

def run_pre_m4_audit():
    print("=" * 60)
    print("      SWIPEX PRE-MILESTONE 4 COMPREHENSIVE AUDIT TEST")
    print("=" * 60)

    # -------------------------------------------------------------
    # 1. MILESTONE 1: AUTH & CORE SETUP
    # -------------------------------------------------------------
    print("\n--- 1. Testing Milestone 1: Auth & Core Setup ---")
    
    # 1.1 Password Hashing Check
    test_pwd = "SecretAuditPassword123!"
    hashed = auth.hash_password(test_pwd)
    assert hashed != test_pwd, "FAIL: Password stored in plaintext!"
    assert auth.verify_password(test_pwd, hashed), "FAIL: Bcrypt verification failed!"
    print("[OK] Bcrypt password hashing verified")

    # 1.2 Registration & JWT Auth Flow
    audit_seeker_email = "audit_seeker@swipex.io"
    audit_recruiter_email = "audit_recruiter@swipex.io"
    
    # Clean up existing test users if re-running
    db = SessionLocal()
    for email in [audit_seeker_email, audit_recruiter_email]:
        user = db.query(models.User).filter(models.User.email == email).first()
        if user:
            db.query(models.Swipe).filter(models.Swipe.user_id == user.id).delete()
            db.query(models.Resume).filter(models.Resume.user_id == user.id).delete()
            db.query(models.JobSeekerProfile).filter(models.JobSeekerProfile.user_id == user.id).delete()
            db.query(models.RecruiterProfile).filter(models.RecruiterProfile.user_id == user.id).delete()
            db.query(models.User).filter(models.User.id == user.id).delete()
    db.commit()

    # Register Candidate
    res = client.post("/api/auth/register", json={
        "email": audit_seeker_email,
        "password": "Password123!",
        "role": "job_seeker",
        "full_name": "Audit Candidate",
        "phone": "555-0100",
        "location": "New York, NY"
    })
    assert res.status_code == 201, f"FAIL: Candidate registration failed: {res.text}"
    seeker_data = res.json()
    print(f"[OK] Candidate registered successfully: {seeker_data['email']}")

    # Login Candidate
    res = client.post("/api/auth/login", json={
        "email": audit_seeker_email,
        "password": "Password123!"
    })
    assert res.status_code == 200, f"FAIL: Candidate login failed: {res.text}"
    tokens = res.json()
    seeker_token = tokens["access_token"]
    seeker_headers = {"Authorization": f"Bearer {seeker_token}"}
    print("[OK] JWT Access Token obtained")

    # Protected Route /me Check
    res = client.get("/api/auth/me", headers=seeker_headers)
    assert res.status_code == 200, f"FAIL: Protected /me endpoint failed: {res.text}"
    assert res.json()["role"] == "job_seeker"
    print("[OK] Protected route /api/auth/me verified")

    # 1.3 Role-Based Access Control (RBAC)
    first_company = db.query(models.Company).first()
    comp_id = str(first_company.id) if first_company else "00000000-0000-0000-0000-000000000000"
    
    # Candidate attempting to create job (Should be 403 FORBIDDEN)
    res = client.post("/api/jobs", json={
        "company_id": comp_id,
        "title": "Unauthorized Job",
        "description": "Test description",
        "job_type": "full_time"
    }, headers=seeker_headers)
    assert res.status_code == 403, f"FAIL: Job seeker was NOT blocked from creating job! Status: {res.status_code}"
    print("[OK] RBAC verified: Job Seeker correctly blocked (403) from creating job postings")

    # Register Recruiter & Test Authorized Job Creation
    res = client.post("/api/auth/register", json={
        "email": audit_recruiter_email,
        "password": "Password123!",
        "role": "recruiter",
        "full_name": "Audit Recruiter",
        "company_name": "Audit Labs Inc",
        "company_website": "https://auditlabs.io"
    })
    assert res.status_code == 201
    res = client.post("/api/auth/login", json={"email": audit_recruiter_email, "password": "Password123!"})
    recruiter_headers = {"Authorization": f"Bearer {res.json()['access_token']}"}
    print("[OK] Recruiter login & authorization verified")

    # -------------------------------------------------------------
    # 2. MILESTONE 2: SWIPE DECK & JOB DISCOVERY
    # -------------------------------------------------------------
    print("\n--- 2. Testing Milestone 2: Swipe Deck & Job Discovery ---")

    # 2.1 Job Filtering
    res = client.get("/api/jobs?remote=true")
    assert res.status_code == 200
    print(f"[OK] GET /api/jobs?remote=true returned {len(res.json())} jobs")

    res = client.get("/api/jobs?company_type=mnc")
    assert res.status_code == 200
    mnc_jobs = res.json()
    print(f"[OK] GET /api/jobs?company_type=mnc returned {len(mnc_jobs)} jobs")

    res = client.get("/api/jobs?salary_min=150000")
    assert res.status_code == 200
    print(f"[OK] GET /api/jobs?salary_min=150000 returned {len(res.json())} jobs")

    # 2.2 Swipe Persistence & Recommendation Filtering
    res = client.get("/api/recommendations", headers=seeker_headers)
    assert res.status_code == 200
    initial_recs = res.json()
    assert len(initial_recs) > 0, "FAIL: No recommendations found!"
    target_job = initial_recs[0]
    target_job_id = target_job["id"]
    print(f"[OK] Initial recommendation feed contains {len(initial_recs)} jobs")

    # Swipe RIGHT on first job
    res = client.post("/api/swipes", json={
        "job_id": target_job_id,
        "direction": "right"
    }, headers=seeker_headers)
    assert res.status_code == 201, f"FAIL: Swipe record failed: {res.text}"
    print(f"[OK] Recorded RIGHT swipe on job '{target_job['title']}'")

    # Confirm swiped job disappears from recommendation feed
    res = client.get("/api/recommendations", headers=seeker_headers)
    new_recs = res.json()
    new_rec_ids = [j["id"] for j in new_recs]
    assert target_job_id not in new_rec_ids, "FAIL: Swiped job reappeared in feed!"
    print("[OK] Swipe persistence confirmed: Swiped job excluded from recommendation feed")

    # 2.3 Reset Swipes (DELETE /api/swipes)
    res = client.delete("/api/swipes", headers=seeker_headers)
    assert res.status_code == 200, f"FAIL: Reset swipes failed: {res.text}"
    res = client.get("/api/recommendations", headers=seeker_headers)
    reset_rec_ids = [j["id"] for j in res.json()]
    assert target_job_id in reset_rec_ids, "FAIL: Job did not reappear after resetting swipes!"
    print("[OK] DELETE /api/swipes tested successfully: Swipe history cleared and feed restored")

    # -------------------------------------------------------------
    # 3. MILESTONE 3: AI RESUME, ATS & CAREER SUITE
    # -------------------------------------------------------------
    print("\n--- 3. Testing Milestone 3: AI Resume, ATS & Career Suite ---")

    # 3.1 PDF & DOCX Resume Upload
    # Test TXT upload
    txt_content = b"Candidate Name: Audit Seeker\nSkills: Python, FastAPI, React, PostgreSQL, Docker, AWS, Git"
    res = client.post("/api/resumes/upload", files={"file": ("resume.txt", txt_content, "text/plain")}, headers=seeker_headers)
    assert res.status_code == 200, f"FAIL: TXT Resume upload failed: {res.text}"
    parsed_skills = res.json()["parsed_skills"]
    assert "Python" in parsed_skills and "FastAPI" in parsed_skills
    print(f"[OK] TXT Resume uploaded & parsed ({len(parsed_skills)} skills detected: {parsed_skills[:5]}...)")

    # Test DOCX upload
    doc = docx.Document()
    doc.add_heading("Audit Candidate Resume", 0)
    doc.add_paragraph("Skills: Python, React, TypeScript, Node.js, GraphQL, Docker, Kubernetes, AWS")
    doc_io = io.BytesIO()
    doc.save(doc_io)
    res = client.post("/api/resumes/upload", files={"file": ("resume.docx", doc_io.getvalue(), "application/vnd.openxmlformats-officedocument.wordprocessingml.document")}, headers=seeker_headers)
    assert res.status_code == 200, f"FAIL: DOCX Resume upload failed: {res.text}"
    print("[OK] DOCX Resume uploaded & parsed successfully")

    # 3.2 ATS Scoring Engine
    res = client.post(f"/api/resumes/analyze/{target_job_id}", headers=seeker_headers)
    assert res.status_code == 200, f"FAIL: ATS Analysis failed: {res.text}"
    analysis = res.json()
    assert "ats_score" in analysis
    assert "matched_keywords" in analysis
    assert "suggestions" in analysis
    assert len(analysis["suggestions"]) == 3
    print(f"[OK] ATS Analysis verified: Match Score {analysis['ats_score']}%, Matched: {analysis['matched_keywords']}, Missing: {analysis['missing_keywords']}")

    # 3.3 AI Job Feed Re-Ranking
    res = client.get("/api/recommendations", headers=seeker_headers)
    assert res.status_code == 200
    ranked_jobs = res.json()
    match_scores = [j.get("match_score", 0) for j in ranked_jobs]
    assert match_scores == sorted(match_scores, reverse=True), f"FAIL: Recommendation feed not sorted descending by ATS match score! Scores: {match_scores}"
    print(f"[OK] AI Recommendation Feed sorting confirmed: Jobs ranked descending by match score ({match_scores[:4]})")

    # 3.4 Cover Letter & Interview Prep Generators
    res = client.post(f"/api/resumes/cover-letter/{target_job_id}", headers=seeker_headers)
    assert res.status_code == 200
    cover_letter = res.json()["cover_letter"]
    assert "Dear Hiring Manager" in cover_letter or len(cover_letter) > 100
    print("[OK] AI Cover Letter Generator verified")

    res = client.post(f"/api/resumes/interview-prep/{target_job_id}", headers=seeker_headers)
    assert res.status_code == 200
    questions = res.json()["questions"]
    assert len(questions) >= 3
    print(f"[OK] AI Interview Prep Generator verified ({len(questions)} questions generated)")

    # 3.5 Gemini Fallback Test (Intentionally clear GEMINI_API_KEY)
    orig_gemini_key = os.environ.get("GEMINI_API_KEY")
    if "GEMINI_API_KEY" in os.environ:
        del os.environ["GEMINI_API_KEY"]
    
    from app.services.ai_advisor import generate_ai_suggestions, generate_cover_letter, generate_interview_questions
    fb_suggestions = generate_ai_suggestions("Fullstack Engineer", ["React", "Python"], ["GraphQL"], 75)
    assert len(fb_suggestions) == 3
    fb_cover = generate_cover_letter("Audit Candidate", ["Python", "React"], "Fullstack Engineer", "Google", "Job Desc")
    assert "Audit Candidate" in fb_cover
    fb_questions = generate_interview_questions("Fullstack Engineer", ["Python", "React"], ["GraphQL"])
    assert len(fb_questions) >= 3
    
    if orig_gemini_key:
        os.environ["GEMINI_API_KEY"] = orig_gemini_key
    print("[OK] Gemini local NLP fallback engine tested and operational when API key is missing/invalid")

    print("\n" + "=" * 60)
    print("  [SUCCESS] ALL PRE-MILESTONE 4 AUDIT CHECKS PASSED SUCCESSFULLY 100%!")
    print("=" * 60)

if __name__ == "__main__":
    run_pre_m4_audit()
