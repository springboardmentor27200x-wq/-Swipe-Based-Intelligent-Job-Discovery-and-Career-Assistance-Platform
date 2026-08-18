"""
SwipeX Resumes — Pytest test suite (Milestone 3)
Tests: parsing, ATS scoring, upload API, resume management, job-compatibility API,
       and the resume-aware recommendation engine.
"""

import io
import pytest
from rest_framework.test import APIClient
from rest_framework import status as http_status

from apps.users.models import User
from apps.jobs.models import Company, Job, Skill
from apps.resumes.models import Resume, ATSScore
from apps.resumes.parsing import parse_resume_text, extract_text_and_links
from apps.resumes.scoring import compute_ats_score
from apps.jobs.services import compute_recommendation_score


SAMPLE_RESUME_TEXT = """
Pratyusha Satpathy
pratyusha@example.com | +91 9876543210
linkedin.com/in/pratyusha | github.com/pratyusha

SKILLS
Python, Django, React, PostgreSQL, Docker, REST API, Git

EXPERIENCE
Software Engineer Intern, Acme Corp (2022 - 2024)
Built REST APIs using Django and deployed with Docker on AWS.

PROJECTS
SwipeX - A swipe based job discovery platform built with Django and React.

EDUCATION
Bachelor of Technology in Computer Science, XYZ University (2020 - 2024)

CERTIFICATIONS
AWS Certified Cloud Practitioner
"""


def _build_docx_bytes(text_lines):
    import docx
    document = docx.Document()
    for line in text_lines:
        document.add_paragraph(line)
    buf = io.BytesIO()
    document.save(buf)
    buf.seek(0)
    return buf.read()


def _build_pdf_bytes(text_lines):
    from reportlab.pdfgen import canvas
    buf = io.BytesIO()
    c = canvas.Canvas(buf)
    y = 800
    for line in text_lines:
        c.drawString(50, y, line)
        y -= 18
        if y < 50:
            c.showPage()
            y = 800
    c.save()
    buf.seek(0)
    return buf.read()


# ── Fixtures ──────────────────────────────────────────────────────────────────

@pytest.fixture
def api():
    return APIClient()


@pytest.fixture
def seeker(db):
    return User.objects.create_user(
        email='seeker@test.com', password='Pass1234!',
        role=User.Role.JOB_SEEKER, first_name='Ana', last_name='Seeker',
        is_email_verified=True,
    )


@pytest.fixture
def recruiter(db):
    return User.objects.create_user(
        email='recruiter@test.com', password='Pass1234!',
        role=User.Role.RECRUITER, first_name='Raj', last_name='Recruiter',
    )


@pytest.fixture
def company(db, recruiter):
    return Company.objects.create(name='AcmeCorp', company_type='startup', industry='Technology', recruiter=recruiter)


@pytest.fixture
def skills(db):
    names = ['Python', 'Django', 'React', 'PostgreSQL', 'Docker', 'REST API', 'Git', 'AWS', 'Kubernetes']
    objs = {}
    for n in names:
        objs[n] = Skill.objects.create(name=n)
    return objs


@pytest.fixture
def published_job(db, recruiter, company, skills):
    job = Job.objects.create(
        recruiter=recruiter, company=company,
        title='Backend Engineer', description='Build great REST APIs using Django and Docker on AWS.',
        requirements='Experience with Python, Django, Docker, AWS, Kubernetes required.',
        experience_level='junior', job_type='full_time',
        work_mode='remote', location='Bangalore', status='published',
    )
    job.skills_required.set([skills['Python'], skills['Django'], skills['Docker'], skills['AWS'], skills['Kubernetes']])
    from django.utils import timezone
    job.published_at = timezone.now()
    job.save()
    return job


def _auth(api, user):
    api.force_authenticate(user=user)
    return api


# ══════════════════════════════════════════════════════════════════════════════
# PARSING
# ══════════════════════════════════════════════════════════════════════════════

class TestResumeParsing:
    def test_parses_skills_education_experience_projects_certifications(self, db):
        result = parse_resume_text(SAMPLE_RESUME_TEXT)
        assert 'Python' in result['parsed_skills']
        assert 'Django' in result['parsed_skills']
        assert 'React' in result['parsed_skills']
        assert result['parsed_education']
        assert result['parsed_experience']
        assert result['parsed_projects']
        assert result['parsed_certifications']
        assert result['parsed_email'] == 'pratyusha@example.com'
        assert result['has_github_link'] is True
        assert result['has_linkedin_link'] is True
        assert result['estimated_years_experience'] > 0

    def test_handles_empty_text_gracefully(self, db):
        result = parse_resume_text("")
        assert result['parsed_skills'] == []
        assert result['parsed_education'] == []


# ══════════════════════════════════════════════════════════════════════════════
# ATS SCORING
# ══════════════════════════════════════════════════════════════════════════════

class TestATSScoring:
    def test_full_skill_match_scores_high(self, db, seeker, published_job):
        resume = Resume.objects.create(
            user=seeker, file=b'', original_filename='r.pdf', file_type='pdf',
            raw_text=SAMPLE_RESUME_TEXT,
            parsed_skills=['Python', 'Django', 'Docker', 'AWS', 'Kubernetes'],
            parsed_technologies=['Python', 'Django', 'Docker', 'AWS', 'Kubernetes', 'React'],
            parsed_education=['Bachelor of Technology in Computer Science, XYZ University (2020 - 2024)'],
            parsed_experience=['Software Engineer Intern, Acme Corp (2022 - 2024)'],
            parsed_projects=['SwipeX project'],
            parsed_certifications=['AWS Certified Cloud Practitioner'],
            has_github_link=True, has_linkedin_link=True,
            estimated_years_experience=2,
            parse_status=Resume.ParseStatus.SUCCESS,
        )
        result = compute_ats_score(resume, published_job)
        assert result['overall_score'] > 70
        assert result['skill_match'] == 100.0
        assert result['missing_skills'] == []

    def test_missing_skills_detected_and_suggested(self, db, seeker, published_job):
        resume = Resume.objects.create(
            user=seeker, file=b'', original_filename='r.pdf', file_type='pdf',
            raw_text='Just some generic resume text without much detail.',
            parsed_skills=['Python'],
            parsed_technologies=['Python'],
            parsed_education=[], parsed_experience=[], parsed_projects=[], parsed_certifications=[],
            has_github_link=False, has_linkedin_link=False,
            estimated_years_experience=0,
            parse_status=Resume.ParseStatus.SUCCESS,
        )
        result = compute_ats_score(resume, published_job)
        assert 'Django' in result['missing_skills']
        assert 'Docker' in result['missing_skills']
        assert result['overall_score'] < 70
        assert len(result['suggestions']) > 0
        assert any('Django' in s or 'Docker' in s or 'AWS' in s or 'Kubernetes' in s for s in result['suggestions'])

    def test_compatibility_label_thresholds(self, db, seeker, published_job):
        ats = ATSScore(overall_score=90)
        assert ats.compatibility_label == 'Excellent'
        ats.overall_score = 75
        assert ats.compatibility_label == 'Good'
        ats.overall_score = 55
        assert ats.compatibility_label == 'Fair'
        ats.overall_score = 20
        assert ats.compatibility_label == 'Poor'


# ══════════════════════════════════════════════════════════════════════════════
# UPLOAD / MANAGEMENT API
# ══════════════════════════════════════════════════════════════════════════════

class TestResumeUploadAPI:
    def test_upload_docx_resume_parses_successfully(self, api, seeker):
        _auth(api, seeker)
        content = _build_docx_bytes(SAMPLE_RESUME_TEXT.strip().split('\n'))
        from django.core.files.uploadedfile import SimpleUploadedFile
        f = SimpleUploadedFile('resume.docx', content,
                                content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        resp = api.post('/api/v1/resumes/upload/', {'file': f}, format='multipart')
        assert resp.status_code == http_status.HTTP_201_CREATED, resp.content
        data = resp.data['data']
        assert data['parse_status'] == 'success'
        assert 'Python' in data['parsed_skills']
        assert data['is_primary'] is True

    def test_upload_pdf_resume_parses_successfully(self, api, seeker):
        _auth(api, seeker)
        content = _build_pdf_bytes(SAMPLE_RESUME_TEXT.strip().split('\n'))
        from django.core.files.uploadedfile import SimpleUploadedFile
        f = SimpleUploadedFile('resume.pdf', content, content_type='application/pdf')
        resp = api.post('/api/v1/resumes/upload/', {'file': f}, format='multipart')
        assert resp.status_code == http_status.HTTP_201_CREATED, resp.content
        data = resp.data['data']
        assert data['parse_status'] == 'success'
        assert 'Python' in data['parsed_skills']

    def test_rejects_unsupported_file_type(self, api, seeker):
        _auth(api, seeker)
        from django.core.files.uploadedfile import SimpleUploadedFile
        f = SimpleUploadedFile('resume.txt', b'hello', content_type='text/plain')
        resp = api.post('/api/v1/resumes/upload/', {'file': f}, format='multipart')
        assert resp.status_code == http_status.HTTP_400_BAD_REQUEST

    def test_second_upload_becomes_primary_and_first_demoted(self, api, seeker):
        _auth(api, seeker)
        from django.core.files.uploadedfile import SimpleUploadedFile
        for i in range(2):
            content = _build_docx_bytes(SAMPLE_RESUME_TEXT.strip().split('\n'))
            f = SimpleUploadedFile(f'resume{i}.docx', content,
                                    content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
            api.post('/api/v1/resumes/upload/', {'file': f}, format='multipart')

        assert Resume.objects.filter(user=seeker, is_primary=True).count() == 1
        assert Resume.objects.filter(user=seeker).count() == 2

    def test_list_and_delete_resume(self, api, seeker):
        resume = Resume.objects.create(
            user=seeker, file='resumes/x/y.pdf', original_filename='r.pdf', file_type='pdf',
            parse_status=Resume.ParseStatus.SUCCESS,
        )
        _auth(api, seeker)
        resp = api.get('/api/v1/resumes/')
        assert resp.status_code == 200
        assert len(resp.data['data']) == 1

        resp = api.delete(f'/api/v1/resumes/{resume.id}/')
        assert resp.status_code == http_status.HTTP_204_NO_CONTENT
        assert Resume.objects.filter(id=resume.id).count() == 0

    def test_only_job_seekers_can_upload(self, api, recruiter):
        _auth(api, recruiter)
        from django.core.files.uploadedfile import SimpleUploadedFile
        f = SimpleUploadedFile('resume.docx', b'x', content_type='text/plain')
        resp = api.post('/api/v1/resumes/upload/', {'file': f}, format='multipart')
        assert resp.status_code == http_status.HTTP_403_FORBIDDEN


# ══════════════════════════════════════════════════════════════════════════════
# JOB COMPATIBILITY API
# ══════════════════════════════════════════════════════════════════════════════

class TestResumeJobMatchAPI:
    def test_match_endpoint_without_resume_returns_none(self, api, seeker, published_job):
        _auth(api, seeker)
        resp = api.get(f'/api/v1/resumes/match/{published_job.id}/')
        assert resp.status_code == 200
        assert resp.data['data'] is None

    def test_match_endpoint_with_resume_returns_scores(self, api, seeker, published_job):
        Resume.objects.create(
            user=seeker, file='resumes/x/y.pdf', original_filename='r.pdf', file_type='pdf',
            raw_text=SAMPLE_RESUME_TEXT,
            parsed_skills=['Python', 'Django', 'Docker'],
            parsed_technologies=['Python', 'Django', 'Docker'],
            estimated_years_experience=2,
            parse_status=Resume.ParseStatus.SUCCESS,
        )
        _auth(api, seeker)
        resp = api.get(f'/api/v1/resumes/match/{published_job.id}/')
        assert resp.status_code == 200
        data = resp.data['data']
        assert 'overall_score' in data
        assert 'compatibility_label' in data
        assert 'missing_skills' in data
        assert 'suggestions' in data

    def test_job_detail_api_includes_ats_fields_for_seeker_with_resume(self, api, seeker, published_job):
        Resume.objects.create(
            user=seeker, file='resumes/x/y.pdf', original_filename='r.pdf', file_type='pdf',
            raw_text=SAMPLE_RESUME_TEXT,
            parsed_skills=['Python', 'Django'],
            parsed_technologies=['Python', 'Django'],
            parse_status=Resume.ParseStatus.SUCCESS,
        )
        _auth(api, seeker)
        resp = api.get(f'/api/v1/jobs/{published_job.id}/')
        assert resp.status_code == 200
        data = resp.data['data']
        assert data['ats_score'] is not None
        assert data['ats_compatibility'] in ['Excellent', 'Good', 'Fair', 'Poor']
        assert isinstance(data['missing_skills'], list)
        assert isinstance(data['resume_suggestions'], list)

    def test_job_list_api_includes_ats_fields(self, api, seeker, published_job):
        Resume.objects.create(
            user=seeker, file='resumes/x/y.pdf', original_filename='r.pdf', file_type='pdf',
            parsed_skills=['Python'], parsed_technologies=['Python'],
            parse_status=Resume.ParseStatus.SUCCESS,
        )
        _auth(api, seeker)
        resp = api.get('/api/v1/jobs/')
        assert resp.status_code == 200
        job_data = next(j for j in resp.data['data'] if j['id'] == str(published_job.id))
        assert 'ats_score' in job_data
        assert 'ats_compatibility' in job_data


# ══════════════════════════════════════════════════════════════════════════════
# RECOMMENDATION ENGINE — RESUME AWARENESS
# ══════════════════════════════════════════════════════════════════════════════

class TestResumeAwareRecommendations:
    def test_resume_skills_boost_recommendation_score(self, db, seeker, published_job):
        score_without_resume, _ = compute_recommendation_score(seeker, published_job)

        Resume.objects.create(
            user=seeker, file='resumes/x/y.pdf', original_filename='r.pdf', file_type='pdf',
            parsed_skills=['Python', 'Django', 'Docker', 'AWS', 'Kubernetes'],
            parsed_technologies=['Python', 'Django', 'Docker', 'AWS', 'Kubernetes'],
            parse_status=Resume.ParseStatus.SUCCESS,
        )
        score_with_resume, reasons_with_resume = compute_recommendation_score(seeker, published_job)

        assert score_with_resume >= score_without_resume
        assert any('Resume matches' in r or 'Matching skills' in r for r in reasons_with_resume)

    def test_recommendation_engine_still_works_without_any_resume(self, db, seeker, published_job):
        score, reasons = compute_recommendation_score(seeker, published_job)
        assert 0.0 <= score <= 1.0

    def test_semantic_similarity_factor_included_when_resume_has_text(self, db, seeker, published_job):
        """
        Milestone 3.1: resume text closely matching the job description should
        surface a semantic-match reason, on top of (not instead of) the
        existing rule-based skill/location/etc. factors.
        """
        published_job.description = 'We need a Python Django backend engineer with Docker and AWS experience.'
        published_job.requirements = 'Strong Python, Django, Docker, AWS, Kubernetes background required.'
        published_job.save()

        Resume.objects.create(
            user=seeker, file='resumes/x/y.pdf', original_filename='r.pdf', file_type='pdf',
            raw_text='Backend engineer experienced in Python, Django, Docker, AWS and Kubernetes deployments.',
            parsed_skills=['Python', 'Django', 'Docker'],
            parsed_technologies=['Python', 'Django', 'Docker'],
            parse_status=Resume.ParseStatus.SUCCESS,
        )
        score, reasons = compute_recommendation_score(seeker, published_job)
        assert 0.0 <= score <= 1.0
        assert any('semantic' in r.lower() for r in reasons)


# ══════════════════════════════════════════════════════════════════════════════
# PROJECTS PARSING (regression tests — Milestone 3.1 fix)
# ══════════════════════════════════════════════════════════════════════════════

PROJECTS_RESUME_TEXT = """
Pratyusha Satpathy
pratyusha@example.com | +91 9876543210
github.com/pratyusha | linkedin.com/in/pratyusha

TECHNICAL SKILLS
Python, Django, React, PostgreSQL, Git, Docker

EDUCATION
B.Tech in Computer Science, XYZ University, 2020-2024

WORK EXPERIENCE
Software Engineer Intern, ABC Corp, 2023-2024
- Built REST APIs using Django and integrated with React frontend

KEY PROJECTS
SwipeX - Swipe Based Job Discovery Platform
Tech Stack: Django, React, PostgreSQL, Docker
- Built a Tinder-style swipe UI for job discovery
- Implemented JWT authentication and REST APIs
- Deployed using Docker on AWS

Portfolio Website
Technologies: HTML, CSS, JavaScript
- Designed and built a personal portfolio site
- Optimized for performance and SEO

CERTIFICATIONS
AWS Certified Cloud Practitioner
"""


class TestProjectsParsing:
    """
    Regression tests for the Milestone 3.1 fix: the parser used to skip the
    Projects section entirely (it was swallowed by an ambiguous 'Technologies:'
    header match). It must now extract structured project entries.
    """

    def test_projects_section_is_not_skipped(self, db):
        result = parse_resume_text(PROJECTS_RESUME_TEXT)
        assert len(result['parsed_projects']) == 2

    def test_project_title_extracted(self, db):
        result = parse_resume_text(PROJECTS_RESUME_TEXT)
        titles = [p['title'] for p in result['parsed_projects']]
        assert any('SwipeX' in t for t in titles)
        assert any('Portfolio Website' in t for t in titles)

    def test_project_technologies_extracted(self, db):
        result = parse_resume_text(PROJECTS_RESUME_TEXT)
        swipex = next(p for p in result['parsed_projects'] if 'SwipeX' in p['title'])
        assert 'Django' in swipex['technologies']
        assert 'React' in swipex['technologies']
        assert 'PostgreSQL' in swipex['technologies']
        assert 'Docker' in swipex['technologies']

    def test_project_description_extracted(self, db):
        result = parse_resume_text(PROJECTS_RESUME_TEXT)
        swipex = next(p for p in result['parsed_projects'] if 'SwipeX' in p['title'])
        assert 'Tinder-style swipe UI' in swipex['description']

    def test_project_technologies_feed_into_overall_technologies(self, db):
        """ATS scoring must be able to see project-only technologies (e.g. via all_skills)."""
        result = parse_resume_text(PROJECTS_RESUME_TEXT)
        assert 'Docker' in result['parsed_technologies']
        assert 'PostgreSQL' in result['parsed_technologies']

    def test_numbered_project_list_splits_into_separate_entries(self, db):
        text = """
        PROJECTS
        1. E-Commerce Platform
        Built a full e-commerce site using MERN stack.
        * Implemented cart and checkout
        * Used MongoDB for the product catalog

        2. Weather App
        A weather forecasting app using OpenWeather API and React.
        * Responsive UI with Tailwind CSS
        """
        result = parse_resume_text(text)
        assert len(result['parsed_projects']) == 2
        assert 'E-Commerce' in result['parsed_projects'][0]['title']
        assert 'Weather' in result['parsed_projects'][1]['title']

    def test_academic_projects_header_variant_recognized(self, db):
        text = """
        Academic Projects

        Library Management System
        - Developed using Java and MySQL
        - Implemented fine calculation and book reservation modules
        """
        result = parse_resume_text(text)
        assert len(result['parsed_projects']) == 1
        assert 'Library Management System' in result['parsed_projects'][0]['title']
        assert 'Java' in result['parsed_projects'][0]['technologies']
        assert 'MySQL' in result['parsed_projects'][0]['technologies']

    def test_real_world_resume_no_bullet_glyphs_pipe_link_format(self, db):
        """
        Regression test using an actual user-submitted resume where PDF text
        extraction stripped all bullet glyphs, leaving plain description
        sentences with no marker at all — the only project-boundary signal
        is a 'Title | Link tech1, tech2' pattern. Header is also a
        non-trivial variant: 'Projects / Open-Source'.
        """
        text = (
            "Pratyusha Satpathy\n"
            "pratyushasatpathy0@gmail.com | +91-7846802919\n"
            "Linkedin | GitHub\n"
            "Education\n"
            "Odisha University of Technology and Research Bhubaneswar, India\n"
            "Degree in B-Tech- CSE(AIML) Aug -2023 - May-2027\n"
            "CGPA: 9.31\n"
            "Experience\n"
            "IBM SkillsBuild(CSRBOX) | AI Intern - Agentic AI Program Remote | July 2025 - Aug 2025\n"
            "Completed a 4-week project-based learning program on Agentic AI.\n"
            "Skills\n"
            "Programming Languages: C, C++, Python, Java, HTML/CSS, Javascript\n"
            "Libraries/Frameworks: React.js, Node.js, Express.js, Scikit-learn, TensorFlow/Keras\n"
            "Tools / Platforms: Git, GitHub, Google Colab, VS Code, Tableau, Postman\n"
            "Databases: SQL, MongoDB\n"
            "Projects / Open-Source\n"
            "E-Book Management System | Link React.js, Node.js, Express.js, MongoDB, REST APIs\n"
            "Developed a full-stack E-Book Management System using React.js, Node.js, Express.js, and MongoDB\n"
            "to manage books, users, and transactions.\n"
            "Chest X-Ray Pneumonia Detection | Link Python, TensorFlow/Keras, CNN, OpenCV\n"
            "Built a Pneumonia Detection System using CNN, TensorFlow, and Keras.\n"
            "Smart Resume Analyzer | Link Python, Scikit-learn, NLP, Gradio\n"
            "Developed a Smart Resume Analyzer using Python, NLP, TF-IDF, and Cosine Similarity.\n"
            "Certifications\n"
            "Agentic AI: From Learner to Builder - IBM SkillsBuild(CSRBOX-2025)\n"
            "AI Fundamentals - Cisco Network Academy\n"
        )
        result = parse_resume_text(text)

        assert len(result['parsed_projects']) == 3
        titles = [p['title'] for p in result['parsed_projects']]
        assert 'E-Book Management System' in titles
        assert 'Chest X-Ray Pneumonia Detection' in titles
        assert 'Smart Resume Analyzer' in titles

        ebook = next(p for p in result['parsed_projects'] if p['title'] == 'E-Book Management System')
        assert 'React.js' in ebook['technologies']
        assert 'MongoDB' in ebook['technologies']
        assert 'full-stack' in ebook['description']

        # Skills section must also be parsed correctly (not swallowed by the
        # Projects section header ambiguity)
        assert 'Python' in result['parsed_skills']
        assert 'React' in result['parsed_skills'] or 'React.js' in result['parsed_skills']
        assert len(result['parsed_certifications']) == 2


# ══════════════════════════════════════════════════════════════════════════════
# DYNAMIC RESUME SUGGESTIONS (regression tests — Milestone 3.1 fix)
# ══════════════════════════════════════════════════════════════════════════════

class TestDynamicSuggestions:
    """
    Suggestions must be generated from what was actually parsed — never a
    static checklist. A resume with real projects, GitHub, and LinkedIn
    links must NOT be told to add things it already has.
    """

    def test_no_project_suggestion_when_projects_exist(self, db, seeker, published_job):
        resume = Resume.objects.create(
            user=seeker, file='resumes/x/y.pdf', original_filename='r.pdf', file_type='pdf',
            raw_text=PROJECTS_RESUME_TEXT,
            parsed_skills=['Python', 'Django', 'Docker'],
            parsed_technologies=['Python', 'Django', 'Docker', 'React', 'PostgreSQL'],
            parsed_education=['B.Tech in Computer Science, XYZ University, 2020-2024'],
            parsed_experience=['Software Engineer Intern, ABC Corp, 2023-2024'],
            parsed_projects=[
                {'title': 'SwipeX', 'technologies': ['Django', 'React'], 'description': 'A job platform.'}
            ],
            parsed_certifications=['AWS Certified Cloud Practitioner'],
            has_github_link=True, has_linkedin_link=True,
            estimated_years_experience=2,
            parse_status=Resume.ParseStatus.SUCCESS,
        )
        result = compute_ats_score(resume, published_job)
        assert not any('projects' in s.lower() for s in result['suggestions'])
        assert not any('github' in s.lower() for s in result['suggestions'])
        assert not any('linkedin' in s.lower() for s in result['suggestions'])

    def test_project_suggestion_present_when_projects_missing(self, db, seeker, published_job):
        resume = Resume.objects.create(
            user=seeker, file='resumes/x/y.pdf', original_filename='r.pdf', file_type='pdf',
            raw_text='Generic resume with no real sections.',
            parsed_skills=['Python'], parsed_technologies=['Python'],
            parsed_education=[], parsed_experience=[], parsed_projects=[], parsed_certifications=[],
            has_github_link=False, has_linkedin_link=False,
            estimated_years_experience=0,
            parse_status=Resume.ParseStatus.SUCCESS,
        )
        result = compute_ats_score(resume, published_job)
        assert any('projects' in s.lower() for s in result['suggestions'])
        assert any('github' in s.lower() for s in result['suggestions'])
        assert any('linkedin' in s.lower() for s in result['suggestions'])


# ══════════════════════════════════════════════════════════════════════════════
# GITHUB / LINKEDIN DETECTION (regression tests — Milestone 3.2 fix)
# ══════════════════════════════════════════════════════════════════════════════

class TestGithubLinkedinDetection:
    """
    Many resumes render 'GitHub' / 'LinkedIn' as plain clickable anchor text
    with the actual URL only present as an embedded hyperlink (a PDF link
    annotation or a DOCX hyperlink relationship) — not anywhere in the
    visible text layer. The parser must detect these via the `hyperlinks`
    list, not just via regex over the extracted text.
    """

    def test_printed_url_still_detected_via_text(self, db):
        """Regression guard: the original text-regex path must keep working."""
        text = "Jane Doe\ngithub.com/janedoe | linkedin.com/in/janedoe"
        result = parse_resume_text(text)
        assert result['has_github_link'] is True
        assert result['parsed_github_url'] == 'https://github.com/janedoe'
        assert result['has_linkedin_link'] is True
        assert result['parsed_linkedin_url'] == 'https://linkedin.com/in/janedoe'

    def test_no_links_when_absent_from_both_text_and_hyperlinks(self, db):
        text = "Jane Doe\njane@example.com"
        result = parse_resume_text(text, hyperlinks=[])
        assert result['has_github_link'] is False
        assert result['has_linkedin_link'] is False
        assert result['parsed_github_url'] == ''
        assert result['parsed_linkedin_url'] == ''

    def test_embedded_hyperlink_detected_when_text_has_no_url(self, db):
        """
        The exact real-world case: visible text is just 'Linkedin | GitHub'
        with no URL printed anywhere — the destination is only available as
        an embedded hyperlink.
        """
        text = "Pratyusha Satpathy\npratyusha@example.com\nLinkedin | GitHub"
        hyperlinks = [
            'https://linkedin.com/in/pratyusha-satpathy',
            'https://github.com/pratyusha-satpathy',
        ]
        result = parse_resume_text(text, hyperlinks=hyperlinks)
        assert result['has_github_link'] is True
        assert result['parsed_github_url'] == 'https://github.com/pratyusha-satpathy'
        assert result['has_linkedin_link'] is True
        assert result['parsed_linkedin_url'] == 'https://linkedin.com/in/pratyusha-satpathy'

    def test_irrelevant_hyperlinks_ignored(self, db):
        text = "Jane Doe\nPortfolio | Email"
        hyperlinks = ['https://janedoe-portfolio.com', 'mailto:jane@example.com']
        result = parse_resume_text(text, hyperlinks=hyperlinks)
        assert result['has_github_link'] is False
        assert result['has_linkedin_link'] is False

    def test_pdf_link_annotations_extracted(self, db, tmp_path):
        """End-to-end: a real PDF with GitHub/LinkedIn as link annotations
        (no visible URL text) must have both hyperlinks extracted."""
        from reportlab.pdfgen import canvas

        pdf_path = tmp_path / 'linked_resume.pdf'
        c = canvas.Canvas(str(pdf_path))
        c.drawString(50, 800, 'Jane Doe')
        c.drawString(50, 780, 'Linkedin | GitHub')
        c.linkURL('https://linkedin.com/in/janedoe', (50, 775, 95, 790), relative=0)
        c.linkURL('https://github.com/janedoe', (100, 775, 150, 790), relative=0)
        c.save()

        with open(pdf_path, 'rb') as f:
            text, hyperlinks = extract_text_and_links(f, 'pdf')

        assert any('github.com/janedoe' in h for h in hyperlinks)
        assert any('linkedin.com/in/janedoe' in h for h in hyperlinks)

        result = parse_resume_text(text, hyperlinks=hyperlinks)
        assert result['has_github_link'] is True
        assert result['has_linkedin_link'] is True

    def test_docx_hyperlink_relationships_extracted(self, db, tmp_path):
        """End-to-end: a real DOCX with a GitHub hyperlink relationship
        (anchor text 'GitHub', no visible URL) must be detected."""
        import docx
        from docx.oxml.ns import qn
        from docx.oxml.shared import OxmlElement
        from docx.opc.constants import RELATIONSHIP_TYPE

        document = docx.Document()
        document.add_paragraph('Jane Doe')
        p = document.add_paragraph()
        r_id = p.part.relate_to('https://github.com/janedoe', RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
        hyperlink = OxmlElement('w:hyperlink')
        hyperlink.set(qn('r:id'), r_id)
        run = OxmlElement('w:r')
        run.append(OxmlElement('w:rPr'))
        t = OxmlElement('w:t')
        t.text = 'GitHub'
        run.append(t)
        hyperlink.append(run)
        p._p.append(hyperlink)

        docx_path = tmp_path / 'linked_resume.docx'
        document.save(str(docx_path))

        with open(docx_path, 'rb') as f:
            text, hyperlinks = extract_text_and_links(f, 'docx')

        assert any('github.com/janedoe' in h for h in hyperlinks)
        result = parse_resume_text(text, hyperlinks=hyperlinks)
        assert result['has_github_link'] is True
        assert result['parsed_github_url'] == 'https://github.com/janedoe'
