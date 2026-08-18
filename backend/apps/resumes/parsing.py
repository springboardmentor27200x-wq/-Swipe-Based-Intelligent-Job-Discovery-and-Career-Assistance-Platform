"""
SwipeX Resume Parser — Milestone 3

Rule-based, dependency-light resume parsing:
  - Text extraction from PDF (pdfplumber) and DOCX (python-docx)
  - Section detection (Education / Experience / Projects / Certifications / Skills)
  - Skill & technology detection against a curated dictionary + live DB Skill list
  - Name / email / phone / GitHub / LinkedIn extraction via regex heuristics

This intentionally avoids heavyweight NLP models (spaCy/transformers) so the
platform stays fast to install and deploy while remaining fully deterministic
and explainable — appropriate for an ATS-style scoring engine.
"""

import io
import re

# ── Curated technology / skill dictionary ──────────────────────────────────────
# Used in addition to whatever Skill rows already exist in the database, so
# parsing works well even before any jobs/skills have been seeded.
CURATED_SKILLS = [
    # Languages
    'Python', 'Java', 'JavaScript', 'TypeScript', 'C++', 'C#', 'C', 'Go', 'Golang',
    'Rust', 'Ruby', 'PHP', 'Swift', 'Kotlin', 'Scala', 'R', 'MATLAB', 'Perl', 'Dart',
    # Web / Frontend
    'React', 'React.js', 'Redux', 'Next.js', 'Vue', 'Vue.js', 'Angular', 'Svelte',
    'HTML', 'HTML5', 'CSS', 'CSS3', 'Tailwind CSS', 'Bootstrap', 'jQuery', 'Sass',
    'Webpack', 'Vite', 'Framer Motion',
    # Backend / Frameworks
    'Django', 'Django REST Framework', 'FastAPI', 'Flask', 'Node.js', 'Express',
    'Express.js', 'Spring', 'Spring Boot', 'Ruby on Rails', 'Laravel', 'ASP.NET',
    'GraphQL', 'REST API', 'RESTful API', 'gRPC', 'Microservices', 'WebSocket',
    # Databases
    'PostgreSQL', 'MySQL', 'SQLite', 'MongoDB', 'Redis', 'Cassandra', 'DynamoDB',
    'Elasticsearch', 'SQL', 'NoSQL', 'Firebase', 'Oracle', 'MS SQL Server',
    # Cloud / DevOps
    'AWS', 'Azure', 'GCP', 'Google Cloud', 'Docker', 'Kubernetes', 'Terraform',
    'Ansible', 'Jenkins', 'CI/CD', 'GitHub Actions', 'GitLab CI', 'Linux', 'Bash',
    'Nginx', 'Cloudflare', 'Vercel', 'Heroku', 'Render',
    # Data / ML
    'Machine Learning', 'Deep Learning', 'TensorFlow', 'PyTorch', 'Keras',
    'scikit-learn', 'Pandas', 'NumPy', 'Data Science', 'Data Analysis',
    'Natural Language Processing', 'NLP', 'Computer Vision', 'OpenCV',
    'sentence-transformers', 'spaCy', 'OpenAI API', 'LLM', 'Power BI', 'Tableau',
    # Tools / Practices
    'Git', 'GitHub', 'GitLab', 'Bitbucket', 'Jira', 'Postman', 'Figma',
    'Agile', 'Scrum', 'TDD', 'Unit Testing', 'Pytest', 'Jest', 'Selenium',
    'React Testing Library', 'OAuth2', 'JWT Authentication', 'Redux / Context API',
    # Mobile
    'React Native', 'Flutter', 'Android', 'iOS',
]

SECTION_HEADERS = {
    'education':      [
        r'education', r'academic background', r'qualifications',
        r'educational qualifications?', r'academic qualifications?',
        r'education\s*&\s*qualifications?',
    ],
    'experience':      [
        r'experience', r'work experience', r'employment history',
        r'professional experience', r'work history', r'career history',
        r'internships?', r'internship experience', r'relevant experience',
    ],
    'projects':        [
        r'projects?', r'personal projects?', r'academic projects?',
        r'key projects?', r'major projects?', r'notable projects?',
        r'selected projects?', r'project experience', r'projects?\s*&\s*publications?',
        r'academic (and|&) personal projects?', r'personal (and|&) academic projects?',
        r'technical projects?', r'mini projects?', r'projects? undertaken',
        r'projects?\s*(/|&|and)\s*open[\s-]?source(\s*(projects?|contributions?))?',
        r'open[\s-]?source(\s*(/|&|and)\s*projects?)?(\s*contributions?)?',
    ],
    'certifications':  [
        r'certifications?', r'licenses?( and certifications?)?',
        r'courses?', r'courses? (and|&) certifications?',
        r'certifications?\s*(and|&)\s*courses?', r'licenses?\s*(and|&)\s*certifications?',
        r'training (and|&) certifications?',
    ],
    'skills':          [
        r'skills', r'technical skills', r'technologies', r'core competencies',
        r'skills\s*(and|&)\s*tools', r'tools\s*(and|&)\s*technologies',
        r'key skills', r'skill set', r'areas of expertise',
    ],
}

EMAIL_RE       = re.compile(r'[\w.+-]+@[\w-]+\.[\w.-]+')
PHONE_CANDIDATE_RE = re.compile(r'\+?\d[\d\s\-().]{7,}\d')
GITHUB_RE      = re.compile(r'(github\.com/[\w\-.]+)', re.IGNORECASE)
LINKEDIN_RE    = re.compile(r'(linkedin\.com/in/[\w\-.]+)', re.IGNORECASE)


def _extract_phone(text: str) -> str:
    """
    Finds the first plausible phone number: a run of separators/digits whose
    digit-only length is 10-13 (avoids false positives like '2022 - 2024').
    """
    for match in PHONE_CANDIDATE_RE.finditer(text):
        candidate = match.group(0)
        digits_only = re.sub(r'\D', '', candidate)
        if 10 <= len(digits_only) <= 13:
            return candidate.strip()
    return ''


class ResumeParseError(Exception):
    pass


def extract_text(file_obj, file_type: str) -> str:
    """Extract raw text from an uploaded PDF or DOCX file object (text only)."""
    text, _links = extract_text_and_links(file_obj, file_type)
    return text


def extract_text_and_links(file_obj, file_type: str):
    """
    Extract raw text AND embedded hyperlink URLs from an uploaded PDF/DOCX.

    Many resumes render GitHub/LinkedIn as plain clickable text ("GitHub",
    "LinkedIn") with the actual URL only present as a PDF link annotation or
    a DOCX hyperlink relationship — not in the visible text layer at all.
    Returning these separately lets the parser detect such links reliably.

    Returns: (text: str, hyperlinks: list[str])
    """
    file_type = (file_type or '').lower()
    try:
        if file_type == 'pdf':
            return _extract_pdf_text_and_links(file_obj)
        elif file_type == 'docx':
            return _extract_docx_text_and_links(file_obj)
        else:
            raise ResumeParseError(f'Unsupported file type: {file_type}')
    except ResumeParseError:
        raise
    except Exception as exc:  # pragma: no cover - defensive
        raise ResumeParseError(f'Failed to read {file_type.upper()} file: {exc}')


def _extract_pdf_text_and_links(file_obj):
    import pdfplumber
    text_parts = []
    hyperlinks = []
    file_obj.seek(0)
    with pdfplumber.open(file_obj) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text() or ''
            text_parts.append(page_text)
            try:
                for link in (page.hyperlinks or []):
                    uri = link.get('uri')
                    if uri:
                        hyperlinks.append(uri)
            except Exception:
                pass  # some PDFs have malformed annotation dicts — text extraction still succeeds
    text = '\n'.join(text_parts).strip()
    if not text:
        raise ResumeParseError('No extractable text found in PDF (it may be scanned/image-based).')
    return text, hyperlinks


def _extract_docx_text_and_links(file_obj):
    import docx
    file_obj.seek(0)
    document = docx.Document(file_obj)
    parts = [p.text for p in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text)
    text = '\n'.join(p for p in parts if p and p.strip())
    if not text.strip():
        raise ResumeParseError('No extractable text found in DOCX.')

    hyperlinks = []
    try:
        for rel in document.part.rels.values():
            if 'hyperlink' in rel.reltype and rel.is_external:
                hyperlinks.append(rel.target_ref)
    except Exception:
        pass  # hyperlink extraction is best-effort — text extraction still succeeds

    return text, hyperlinks


def _normalize_header_candidate(line: str) -> str:
    """Strip bullets/dashes/pipes/underscores/decoration around a candidate header line."""
    s = line.strip()
    s = re.sub(r'^[\-\*\u2022\u25AA\u25CF\u2013\u2014\#\s]+', '', s)
    s = re.sub(r'[\-\*\u2022\u25AA\u25CF\u2013\u2014\#\s:]+$', '', s)
    s = re.sub(r'\s+', ' ', s)
    return s.strip()


_HEADER_PATTERNS = {
    key: [re.compile(r'^(' + pat + r')$', re.IGNORECASE) for pat in pats]
    for key, pats in SECTION_HEADERS.items()
}


def _match_header(line: str):
    """
    If `line` looks like a section header (optionally with inline content
    after a colon, e.g. 'PROJECTS: SwipeX ...'), return (section_key, remainder).
    Otherwise return None.

    Note: inline colon-style headers are NOT matched against the 'skills'
    patterns — a line like 'Technologies: React, Node' inside a Projects
    section is far more often a per-project tech-stack annotation than a
    genuine whole-resume Skills section header (real section headers almost
    always appear alone on their own line).
    """
    candidate = line.strip()
    if not candidate or len(candidate) > 60:
        return None

    head_part, remainder = candidate, ''
    if ':' in candidate:
        before, after = candidate.split(':', 1)
        if len(before.strip()) <= 45:
            head_part, remainder = before, after.strip()

    normalized = _normalize_header_candidate(head_part).lower()
    if not normalized or len(normalized.split()) > 6:
        return None

    for key, patterns in _HEADER_PATTERNS.items():
        if remainder and key == 'skills':
            continue
        for pattern in patterns:
            if pattern.match(normalized):
                return key, remainder
    return None


def _split_sections(text: str) -> dict:
    """Split resume text into named sections based on common ATS headers."""
    lines = text.split('\n')
    sections = {key: [] for key in SECTION_HEADERS}
    sections['_preamble'] = []

    current = '_preamble'
    for line in lines:
        stripped = line.strip()
        if not stripped:
            continue
        match = _match_header(stripped)
        if match:
            current, remainder = match
            if remainder:
                sections[current].append(remainder)
            continue
        sections[current].append(stripped)

    return sections


def _detect_skills(text: str, known_skill_names=None) -> list:
    """Detect known skills/technologies present in the resume text."""
    candidates = set(CURATED_SKILLS)
    if known_skill_names:
        candidates |= set(known_skill_names)

    if not text or not text.strip():
        return []

    text_lower = ' ' + text.lower() + ' '
    found = []
    for skill in candidates:
        skill_lower = skill.lower()
        # word-boundary-ish match; handles skills with special chars like C++, C#, CI/CD
        pattern = r'(?<![\w])' + re.escape(skill_lower) + r'(?![\w])'
        if re.search(pattern, text_lower):
            found.append(skill)

    # Deduplicate case-insensitively, preferring the canonical dictionary casing
    seen = {}
    for skill in found:
        seen.setdefault(skill.lower(), skill)
    return sorted(seen.values())


def _guess_name(text: str) -> str:
    for line in text.split('\n')[:6]:
        candidate = line.strip()
        if not candidate:
            continue
        if EMAIL_RE.search(candidate) or PHONE_CANDIDATE_RE.search(candidate):
            continue
        words = candidate.split()
        if 1 <= len(words) <= 4 and all(w.replace('.', '').isalpha() for w in words):
            return candidate.title()
    return ''


def _estimate_years_experience(experience_lines) -> float:
    """Heuristically estimate total years of experience from date ranges like '2021 - 2023'."""
    years_found = []
    year_pattern = re.compile(r'(19|20)\d{2}')
    for line in experience_lines:
        matches = year_pattern.findall(line)
        full_years = re.findall(r'(19|20)\d{2}', line)
        if len(full_years) >= 1:
            nums = [int(y) for y in re.findall(r'(?:19|20)\d{2}', line)]
            years_found.extend(nums)
    if len(years_found) >= 2:
        return max(0.0, float(max(years_found) - min(years_found)))
    return 0.0


_PROJECT_BULLET_RE = re.compile(r'^[•\-\*▪‣●○◦\u2013\u2014]\s*')
_PROJECT_NUMBERED_RE = re.compile(r'^\d+[.)]\s*')
_PROJECT_TECH_LINE_RE = re.compile(
    r'^(?:tech(?:nolog(?:y|ies))?(?:\s+(?:used|stack))?|tools?(?:\s+used)?|stack|built\s*with|technologies?)'
    r'\s*[:\-]\s*(?P<content>.+)$',
    re.IGNORECASE
)


_PROJECT_LINK_ARTIFACT_RE = re.compile(r'\s*\|\s*(link|github|demo|live demo|source|repo)\b.*$', re.IGNORECASE)


def _clean_project_title(title: str) -> str:
    """Strip trailing '| Link' / '| GitHub' / '| Demo' hyperlink-anchor-text
    artifacts left behind when a PDF's clickable link loses its URL on
    text extraction (e.g. 'MyApp | Link React, Node, MongoDB' -> 'MyApp')."""
    return _PROJECT_LINK_ARTIFACT_RE.sub('', title).strip(' -—|')


_PROJECT_LINK_TITLE_RE = re.compile(
    r'^(?P<title>.+?)\s*\|\s*(?:link|github|demo|live\s*demo|repo|repository|source)\b\s*(?P<techlist>.*)$',
    re.IGNORECASE
)


def _parse_projects(project_lines, known_skill_names=None) -> list:
    """
    Group raw 'Projects' section lines into structured entries:
      { title, technologies: [...], description }

    Heuristics (rule-based, no ML), in priority order:
      1. A "Title | Link tech1, tech2" style line (very common — a hyperlinked
         project title whose URL was stripped by PDF text extraction, leaving
         '| Link' anchor text followed by an inline tech list) is a STRONG,
         unambiguous new-project boundary. Everything up to the next such
         line, numbered marker, tech-line, or bullet is treated as that
         project's description — this matters because many resumes lose
         their bullet glyphs entirely during PDF extraction, leaving no
         other signal to separate projects.
      2. A numbered marker ("1. Project A") is also a strong new-project
         boundary.
      3. A 'Tech Stack:' / 'Technologies:' / 'Built with:' line contributes
         to the current project's technologies.
      4. A symbol-bulleted line (•, -, *, ...) is a description line.
      5. Any other line starts a new project (weak signal) UNLESS the
         current project was opened by a strong signal (in which case it's
         additional description), OR the current project has no content yet
         (in which case it's a continuation of the title/subtitle).
    """
    projects = []
    current = None

    def flush():
        if not current:
            return
        description = ' '.join(current['description_lines']).strip()
        technologies = current['technologies']
        title = _clean_project_title(current['title'])
        if not technologies:
            # Fallback: detect known skills mentioned anywhere in this project's text
            combined = f"{title} {description}"
            technologies = _detect_skills(combined, known_skill_names)
        if title or description or technologies:
            projects.append({
                'title': title,
                'technologies': sorted(set(technologies)),
                'description': description,
            })

    for raw_line in project_lines:
        line = raw_line.strip()
        if not line:
            continue

        link_title_match = _PROJECT_LINK_TITLE_RE.match(line)
        if link_title_match:
            flush()
            techlist_text = link_title_match.group('techlist').strip()
            techs = [t.strip() for t in re.split(r'[,|/]|\band\b', techlist_text) if t.strip()] if techlist_text else []
            current = {
                'title': link_title_match.group('title').strip(),
                'technologies': techs, 'description_lines': [], 'strong_boundary': True,
            }
            continue

        tech_match = _PROJECT_TECH_LINE_RE.match(line)
        if tech_match:
            if current is None:
                current = {'title': '', 'technologies': [], 'description_lines': [], 'strong_boundary': False}
            tokens = re.split(r'[,|/•]|\band\b', tech_match.group('content'))
            current['technologies'].extend(t.strip() for t in tokens if t.strip())
            continue

        # Numbered list markers ("1. Project A") almost always enumerate distinct
        # projects on a resume, so each one unconditionally starts a fresh entry
        # (unlike symbol bullets, which are description points within a project).
        numbered_match = _PROJECT_NUMBERED_RE.match(line)
        if numbered_match:
            flush()
            current = {
                'title': _PROJECT_NUMBERED_RE.sub('', line).strip(),
                'technologies': [], 'description_lines': [], 'strong_boundary': True,
            }
            continue

        bullet_match = _PROJECT_BULLET_RE.match(line)
        if bullet_match:
            desc = _PROJECT_BULLET_RE.sub('', line).strip()
            if current is None:
                current = {'title': '', 'technologies': [], 'description_lines': [], 'strong_boundary': False}
            current['description_lines'].append(desc)
            continue

        # Plain line with no explicit marker.
        if current is not None and current.get('strong_boundary'):
            # This project's boundary was established by a strong signal
            # (link-title or numbered marker) — stay locked onto it until
            # the next strong signal, treating plain lines as description.
            current['description_lines'].append(line)
            continue

        # Weak-signal fallback (no bullets, no strong boundary seen yet):
        # a new project starts once the current one already has content;
        # otherwise this is a continuation of the same title/subtitle.
        if current is None or current['description_lines'] or current['technologies']:
            flush()
            current = {'title': line, 'technologies': [], 'description_lines': [], 'strong_boundary': False}
        else:
            current['title'] = (current['title'] + ' — ' + line).strip(' —')

    flush()
    return projects[:15]


def _find_linked_url(hyperlinks, domain_keywords):
    """Return the first hyperlink URL matching any of the given domain keywords."""
    for url in hyperlinks or []:
        low = (url or '').lower()
        if any(kw in low for kw in domain_keywords):
            return url
    return None


def parse_resume_text(text: str, known_skill_names=None, hyperlinks=None) -> dict:
    """
    Parse extracted resume text into structured fields.
    Returns a dict matching Resume model's parsed_* fields.

    `hyperlinks`: optional list of URLs embedded in the source file as real
    clickable links (PDF link annotations / DOCX hyperlink relationships).
    Many resumes render GitHub/LinkedIn as plain anchor text ("GitHub",
    "LinkedIn") with no visible URL at all — the destination is only
    present as an embedded hyperlink, not in the text layer. Text-based
    regex detection is tried first (it captures the exact printed URL when
    present); embedded hyperlinks are used whenever the text alone doesn't
    contain one.
    """
    sections = _split_sections(text)

    email_match = EMAIL_RE.search(text)
    phone_value = _extract_phone(text)

    github_match = GITHUB_RE.search(text)
    linkedin_match = LINKEDIN_RE.search(text)
    github_url = ('https://' + github_match.group(1)) if github_match else _find_linked_url(hyperlinks, ['github.com'])
    linkedin_url = ('https://' + linkedin_match.group(1)) if linkedin_match else _find_linked_url(hyperlinks, ['linkedin.com'])

    all_detected_skills = _detect_skills(text, known_skill_names)
    skills_section_text = ' '.join(sections.get('skills', []))
    skills_from_section = _detect_skills(skills_section_text, known_skill_names) if skills_section_text else []

    # "skills" = what's explicitly listed in a Skills section (if present),
    # otherwise fall back to whole-document detection.
    parsed_skills = skills_from_section or all_detected_skills

    experience_lines = sections.get('experience', [])
    project_lines = sections.get('projects', [])
    parsed_projects = _parse_projects(project_lines[:80], known_skill_names)

    # "technologies" = everything detected anywhere (union), including
    # technologies mentioned only inside individual project entries.
    project_technologies = {t for proj in parsed_projects for t in proj['technologies']}
    parsed_technologies = sorted(set(all_detected_skills) | project_technologies)

    return {
        'parsed_name': _guess_name(text),
        'parsed_email': email_match.group(0) if email_match else '',
        'parsed_phone': phone_value,
        'parsed_skills': parsed_skills,
        'parsed_technologies': parsed_technologies,
        'parsed_education': sections.get('education', [])[:20],
        'parsed_experience': experience_lines[:30],
        'parsed_projects': parsed_projects,
        'parsed_certifications': sections.get('certifications', [])[:20],
        'has_github_link': bool(github_url),
        'has_linkedin_link': bool(linkedin_url),
        'parsed_github_url': github_url or '',
        'parsed_linkedin_url': linkedin_url or '',
        'estimated_years_experience': _estimate_years_experience(experience_lines),
    }
